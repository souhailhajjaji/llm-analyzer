import re
from pathlib import Path
from typing import Optional
from docx import Document
from dataclasses import dataclass


@dataclass
class ExtractedDocument:
    filename: str
    full_text: str
    paragraphs: list[str]
    tables: list[list[list[str]]]
    metadata: dict


class DocumentExtractor:
    def __init__(self):
        self.max_file_size = 10 * 1024 * 1024

    def extract(self, file_path: Path | str) -> ExtractedDocument:
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.stat().st_size > self.max_file_size:
            raise ValueError(f"File too large: {file_path.stat().st_size} bytes")
        
        if file_path.suffix.lower() != ".docx":
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        doc = Document(str(file_path))
        
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        full_text = self._build_full_text(paragraphs, tables)
        
        metadata = self._extract_metadata(doc, file_path)
        
        return ExtractedDocument(
            filename=file_path.name,
            full_text=full_text,
            paragraphs=paragraphs,
            tables=tables,
            metadata=metadata,
        )

    def _build_full_text(self, paragraphs: list[str], tables: list[list[list[str]]]) -> str:
        parts = []
        
        for para in paragraphs:
            parts.append(para)
        
        for i, table in enumerate(tables):
            if table:
                parts.append(f"\n[Table {i+1}]")
                for row in table:
                    parts.append(" | ".join(row))
        
        return "\n\n".join(parts)

    def _extract_metadata(self, doc, file_path: Path) -> dict:
        metadata = {
            "filename": file_path.name,
            "file_size": file_path.stat().st_size,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        
        core_props = doc.core_properties
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["created"] = str(core_props.created)
        if core_props.modified:
            metadata["modified"] = str(core_props.modified)
        
        return metadata

    def clean_text(self, text: str) -> str:
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        text = re.sub(r'\t+', ' ', text)
        
        text = re.sub(r'\[Page \d+\]', '', text)
        text = re.sub(r'Page \d+ / \d+', '', text)
        
        return text.strip()

    def extract_text_only(self, file_path: Path | str) -> str:
        extracted = self.extract(file_path)
        return self.clean_text(extracted.full_text)
