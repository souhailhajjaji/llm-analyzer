import pytest
import tempfile
from pathlib import Path
from docx import Document

from src.services.document_extractor import DocumentExtractor, ExtractedDocument


@pytest.fixture
def sample_docx():
    doc = Document()
    doc.add_heading('Cahier des Charges - Test', 0)
    doc.add_paragraph('Introduction au projet.')
    doc.add_paragraph('Fonctionnalité principale: Authentification des utilisateurs.')
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Colonne 1'
    table.cell(0, 1).text = 'Colonne 2'
    table.cell(1, 0).text = 'Donnée 1'
    table.cell(1, 1).text = 'Donnée 2'
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        yield tmp.name
    
    Path(tmp.name).unlink(missing_ok=True)


class TestDocumentExtractor:
    def test_extract_success(self, sample_docx):
        extractor = DocumentExtractor()
        result = extractor.extract(sample_docx)
        
        assert isinstance(result, ExtractedDocument)
        assert result.filename.endswith('.docx')
        assert len(result.paragraphs) > 0
        assert len(result.tables) > 0
        assert 'Fonctionnalité principale' in result.full_text
    
    def test_extract_file_not_found(self):
        extractor = DocumentExtractor()
        
        with pytest.raises(FileNotFoundError):
            extractor.extract('/nonexistent/path.docx')
    
    def test_extract_unsupported_type(self):
        extractor = DocumentExtractor()
        
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
            tmp.write(b'test content')
            tmp_path = tmp.name
        
        try:
            with pytest.raises(ValueError):
                extractor.extract(tmp_path)
        finally:
            Path(tmp_path).unlink(missing_ok=True)
    
    def test_clean_text(self):
        extractor = DocumentExtractor()
        
        dirty_text = "Line 1\n\n\n\nLine 2    with spaces"
        cleaned = extractor.clean_text(dirty_text)
        
        assert '\n\n\n\n' not in cleaned
        assert '  ' not in cleaned
    
    def test_extract_text_only(self, sample_docx):
        extractor = DocumentExtractor()
        
        text = extractor.extract_text_only(sample_docx)
        
        assert isinstance(text, str)
        assert len(text) > 0
